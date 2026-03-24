"""
extract_inputs.py — Discovers and extracts text from all IAT input files.

Asset caching strategy:
  - JD, levelling, persona: extracted once → saved to assets/.
    Re-extracted only when source file changes (detected via mtime + size).
  - CVs: always extracted fresh (they change with each batch).

Run from the project working directory. Outputs JSON to stdout.

Usage:  python extract_inputs.py [--force-refresh] [--base-dir PATH]
"""

import argparse
import hashlib
import io
import json
import sys
from datetime import datetime
from pathlib import Path

# Force UTF-8 stdout so Vietnamese / non-ASCII characters don't crash on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

# Resolve base directory: explicit --base-dir arg wins, then CWD.
# Using an arg avoids depending on the shell's CWD, which can drift when
# Claude's Bash tool has executed cd commands earlier in the session.
def _resolve_base_dir() -> Path:
    for i, arg in enumerate(sys.argv[1:], 1):
        if arg == "--base-dir" and i < len(sys.argv):
            return Path(sys.argv[i + 1]).resolve()
        if arg.startswith("--base-dir="):
            return Path(arg.split("=", 1)[1]).resolve()
    return Path.cwd()

BASE_DIR   = _resolve_base_dir()
INPUT_DIR  = BASE_DIR / "input"
CV_DIR     = INPUT_DIR / "cv"
ASSETS_DIR = BASE_DIR / "assets"
MANIFEST   = ASSETS_DIR / ".manifest.json"

EXCLUDE_KEYWORDS = {"JD", "JOB", "LEVEL", "LEVELL", "PERSONA", "INTERVIEWER"}


# ── Low-level extractors ───────────────────────────────────────────────────────

def extract_pdf(path: Path) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def extract_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix == ".docx":
        return extract_docx(path)
    elif suffix in (".xlsx", ".xls"):
        return extract_xlsx(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    return f"[Unsupported format: {suffix}]"


# ── Asset cache helpers ────────────────────────────────────────────────────────

def _file_signature(path: Path) -> str:
    """Cheap change-detection: mtime + size (no full hash needed)."""
    stat = path.stat()
    return f"{stat.st_mtime_ns}:{stat.st_size}"


def _load_manifest() -> dict:
    if MANIFEST.exists():
        try:
            return json.loads(MANIFEST.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _save_manifest(manifest: dict):
    ASSETS_DIR.mkdir(exist_ok=True)
    MANIFEST.write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def _is_fresh(key: str, source_path: Path, manifest: dict) -> bool:
    """Return True if the cached asset matches the current source file."""
    cached_sig = manifest.get(key, {}).get("signature", "")
    return cached_sig == _file_signature(source_path)


def _to_markdown(key: str, raw: str) -> str:
    """Convert raw extracted text to structured Markdown per asset type."""
    if key == "levelling":
        return _levelling_to_md(raw)
    elif key == "jd":
        return _jd_to_md(raw)
    elif key == "persona":
        return _persona_to_md(raw)
    return raw  # fallback: return as-is


def _levelling_to_md(raw: str) -> str:
    """
    Convert the XLSX levelling dump into a Markdown table + per-level detail blocks.

    The xlsx extractor joins each row's cells with ' | ', but multi-line cell
    values (e.g. "Level 2\nMiddle\n(3–5 yrs…)") keep their internal newlines.
    This means a data row may arrive as two consecutive text lines:
        Line A:  "Level 2"
        Line B:  "Middle\n(3–5 yrs…) | CS/IT degree... | Deep LLM..."
    We re-join them before parsing.
    """
    import re

    lines = [l.strip() for l in raw.splitlines()]

    # ── Step 1: group lines into per-level blocks, then flatten ───────────────
    # Each level's data may be spread across multiple lines because the Excel
    # cell "Level 2\nMiddle\n(3–5 yrs…)" becomes 2-3 separate text lines.
    # Strategy: collect all lines between consecutive "Level N" markers into
    # one record, join them with " | " so they form a parseable row.
    level_start = re.compile(r"^Level\s+\d+$", re.IGNORECASE)
    blocks: list[list[str]] = []
    current: list[str] | None = None
    for line in lines:
        if level_start.match(line):
            if current is not None:
                blocks.append(current)
            current = [line]
        elif current is not None:
            current.append(line)
    if current:
        blocks.append(current)

    # Each block: first element is "Level N", remaining elements are either
    # sub-title lines (no pipe) or the data pipe-row.
    # Join the entire block into one string, replacing newlines with " | ".
    joined: list[str] = []
    for block in blocks:
        # Flatten: join all block lines; the pipe-row already has | separators
        flat = " | ".join(part.strip() for part in block if part.strip())
        joined.append(flat)

    # ── Step 2: parse each complete level row ─────────────────────────────────
    # After flattening, each row looks like one of:
    #   Level 0 | Fresher/Intern | <edu> | <know> | <skills> | <exp> | <auto>
    #   Level 1 | Junior | (1-2 yrs / 0-6 mo AI) | <edu> | <know> | <skills> | <exp> | <auto>
    # The year-range sub-line "(N yrs…)" is present for levels 1-4 but absent for level 0.
    year_range_re = re.compile(r"^\(?\d")   # starts with optional "(" then digit

    level_map = {
        "0": ("0–1 yr",  "None"),
        "1": ("1–2 yrs", "0–6 mo"),
        "2": ("3–5 yrs", "1–2 yrs"),
        "3": ("5–8 yrs", "2–4 yrs"),
        "4": ("8+ yrs",  "4+ yrs"),
    }

    records = []
    for line in joined:
        if not re.match(r"Level\s+\d+", line, re.IGNORECASE):
            continue
        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 6:
            continue

        m = re.search(r"\d+", parts[0])
        if not m:
            continue
        n = m.group()

        # parts[1] = title text (e.g. "Fresher / Intern", "Junior")
        title = parts[1] if len(parts) > 1 else f"Level {n}"

        # parts[2] may be the year-range "(1–2 yrs…)" OR the Education column
        offset = 1 if (len(parts) > 2 and year_range_re.match(parts[2])) else 0
        edu    = parts[2 + offset] if len(parts) > 2 + offset else ""
        know   = parts[3 + offset] if len(parts) > 3 + offset else ""
        skills = parts[4 + offset] if len(parts) > 4 + offset else ""
        # experience description & autonomy (last two)
        auto   = parts[-1]

        exp, ai = level_map.get(n, ("?", "?"))
        records.append(dict(n=n, title=title, exp=exp, ai=ai,
                            edu=edu, know=know, skills=skills, auto=auto))

    # ── Step 3: render ────────────────────────────────────────────────────────
    md = ["# Levelling Framework\n",
          "| Level | Title | Total Exp | AI Exp | Key Skills (summary) |",
          "|---|---|---|---|---|"]

    for r in records:
        short = r["skills"][:90].rstrip() + ("…" if len(r["skills"]) > 90 else "")
        md.append(f"| {r['n']} | {r['title']} | {r['exp']} | {r['ai']} | {short} |")

    md.append("")
    for r in records:
        md.append(f"\n## Level {r['n']} — {r['title']}\n")
        md.append(f"**Education:** {r['edu']}\n")
        md.append(f"**Knowledge:** {r['know']}\n")
        md.append(f"**Skills:** {r['skills']}\n")
        md.append(f"**Autonomy:** {r['auto']}\n")

    return "\n".join(md)


def _jd_to_md(raw: str) -> str:
    """Format JD text as structured Markdown."""
    lines = [l.strip() for l in raw.splitlines()]
    md = []
    section_keywords = {
        "job description": "## Responsibilities",
        "qualifications": "## Requirements",
        "qualification": "## Requirements",
        "experience": "## Requirements",
        "benefits": "## Benefits",
        "why itr": "## About the Company",
    }
    current_section = None
    title_written = False

    for line in lines:
        if not line:
            continue
        lower = line.lower()

        # Try to detect the role title (first meaningful non-company line)
        if not title_written and len(line) > 5 and line.isupper():
            md.append(f"# {line.title()}\n")
            title_written = True
            continue

        # Section headers
        matched = next((v for k, v in section_keywords.items() if k in lower), None)
        if matched and len(line) < 60:
            if matched != current_section:
                current_section = matched
                md.append(f"\n{matched}\n")
            continue

        # Bullet points
        if line.startswith("-") or line.startswith("•"):
            md.append(f"- {line.lstrip('-•').strip()}")
        else:
            md.append(line)

    return "\n".join(md)


def _persona_to_md(raw: str) -> str:
    """Format persona description as structured Markdown."""
    lines = [l.strip() for l in raw.splitlines() if l.strip()]
    md = ["# Interviewer Persona\n"]
    for line in lines:
        lower = line.lower()
        if "star" in lower and ("model" in lower or "modal" in lower or "method" in lower):
            md.append(f"\n**Interview Methodology:** STAR (Situation, Task, Action, Result)\n")
        elif any(w in lower for w in ["he is", "she is", "they are"]):
            md.append(f"\n**Style:** {line}\n")
        elif any(w in lower for w in ["years of experience", "experience in"]):
            md.append(f"\n**Background:** {line}\n")
        elif "don't lie" in lower or "truth" in lower or "honest" in lower:
            md.append(f"\n**Values:** {line}\n")
        else:
            md.append(line)
    return "\n".join(md)


def _cache_asset(key: str, source_path: Path, raw_text: str, manifest: dict):
    """Convert raw text to Markdown, write to assets/<key>.md, update manifest."""
    ASSETS_DIR.mkdir(exist_ok=True)
    md_text = _to_markdown(key, raw_text)
    asset_file = ASSETS_DIR / f"{key}.md"
    asset_file.write_text(md_text, encoding="utf-8")
    manifest[key] = {
        "signature":    _file_signature(source_path),
        "source_file":  source_path.name,
        "extracted_at": datetime.now().isoformat(timespec="seconds"),
        "char_count":   len(md_text),
    }


def _read_cached_asset(key: str) -> str | None:
    asset_file = ASSETS_DIR / f"{key}.md"
    if asset_file.exists():
        return asset_file.read_text(encoding="utf-8")
    return None


# ── File discovery ─────────────────────────────────────────────────────────────

def extract_contact(text: str) -> dict:
    """
    Extract email and phone from raw CV text using regex.
    Handles common Vietnamese CV formats including:
      Email: ..., e-mail: ..., Mobile: ..., Phone: ..., Tel: ..., SDT: ...
      (+ 84)..., 84 - xxx..., 0xx xxx xxx, pipe-separated lines.
    """
    import re

    email = ""
    phone = ""

    # ── Email ────────────────────────────────────────────────────────────────
    # Stop at pipe/space/comma/semicolon so adjacent fields on the same line
    # (e.g. "Email:foo@x.com|Mobile:...") are not included in the match.
    email_label = re.search(
        r"(?:e[-\s]?mail|email address)\s*[:\-]\s*([^\s|,;]+@[^\s|,;]+)",
        text, re.IGNORECASE
    )
    email_bare = re.search(r"[\w.+\-]+@[\w.\-]+\.\w{2,}", text)
    if email_label:
        email = email_label.group(1).strip(".,;|")
    elif email_bare:
        email = email_bare.group(0).strip(".,;|")

    # ── Phone ────────────────────────────────────────────────────────────────
    # Labelled: allow optional leading ( for formats like (+84) or (+ 84)
    phone_label = re.search(
        r"(?:mobile|phone|tel|s\.?đ\.?t|hotline|contact|di\s*động)\s*[:\-]?\s*"
        r"([\(\+\d][\d\s\-().+]{7,})",
        text, re.IGNORECASE
    )
    # Bare patterns (no label):
    #   +84...  |  0x xxxxxxxx  |  0xx xxx xxx  |  84 - xxx xxx xxx
    phone_bare = re.search(
        r"(?<!\d)(\+84[\d\s\-]{9,}|0\d[\d\s\-]{8,11}|84[\s\-]+\d[\d\s\-]{7,})",
        text
    )

    def clean_phone(raw: str) -> str:
        # Drop everything once letters start (e.g. " (Zalo)", " ext", etc.)
        raw = re.split(r"[A-Za-zÀ-ỹ]", raw)[0]
        return re.sub(r"[^\d+]", "", raw).strip()

    if phone_label:
        phone = clean_phone(phone_label.group(1))
    elif phone_bare:
        phone = clean_phone(phone_bare.group(1))

    return {"email": email, "phone": phone}


def find_cv_files() -> list[Path]:
    seen, results = set(), []

    def add(p: Path):
        rp = p.resolve()
        if rp not in seen:
            seen.add(rp)
            results.append(p)

    for d in (CV_DIR, INPUT_DIR):
        if d.exists():
            for ext in ("*.pdf", "*.docx"):
                for f in sorted(d.glob(ext)):
                    add(f)

    for ext in ("*.pdf", "*.docx"):
        for f in sorted(BASE_DIR.glob(ext)):
            if not any(kw in f.stem.upper() for kw in EXCLUDE_KEYWORDS):
                add(f)

    return results


def find_one(keywords: list[str], extensions: list[str],
             standard_names: list[str]) -> Path | None:
    for name in standard_names:
        for d in (INPUT_DIR, BASE_DIR):
            p = d / name
            if p.exists():
                return p
    for d in (INPUT_DIR, BASE_DIR):
        if not d.exists():
            continue
        for f in sorted(d.iterdir()):
            if f.is_file() and f.suffix.lower() in extensions:
                if any(kw.upper() in f.stem.upper() for kw in keywords):
                    return f
    return None


# ── Cached asset extraction ────────────────────────────────────────────────────

def get_asset(key: str, path: Path | None, manifest: dict,
              force: bool = False) -> tuple[str, bool]:
    """
    Return (text, was_refreshed).
    Reads from assets/ cache if fresh; otherwise extracts from source and caches.
    """
    if path is None:
        cached = _read_cached_asset(key)
        if cached:
            return cached, False
        return "", False

    if not force and _is_fresh(key, path, manifest):
        cached = _read_cached_asset(key)
        if cached:
            return cached, False

    # Extract and cache
    text = extract_file(path)
    _cache_asset(key, path, text, manifest)
    return text, True


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--force-refresh", action="store_true",
                        help="Re-extract all assets even if cache is fresh")
    parser.add_argument("--base-dir", metavar="PATH",
                        help="Project root directory (defaults to cwd)")
    args = parser.parse_args()

    out = {
        "cv_files": [],
        "jd_text": "",
        "levelling_text": "",
        "persona_text": "",
        "hm_clarification_text": "",
        "assets_refreshed": [],
        "errors": [],
    }

    manifest = _load_manifest()

    # ── Stable assets (cached) ────────────────────────────────────────────────
    asset_specs = {
        "jd":        dict(keywords=["JD","JOB","JOBDESC"],
                          extensions=[".txt",".docx",".pdf"],
                          standard_names=["jd.txt","jd.docx"]),
        "levelling": dict(keywords=["LEVEL","LEVELL","LEVELING","LEVELLING"],
                          extensions=[".xlsx",".docx"],
                          standard_names=["levelling.xlsx","leveling.xlsx"]),
        "persona":   dict(keywords=["PERSONA","INTERVIEWER"],
                          extensions=[".docx",".xlsx"],
                          standard_names=["personas.docx","personas.xlsx"]),
        "hm_clarification": dict(keywords=["HM","CLARIF","HIRING_MANAGER","HM_CANDIDATE"],
                                 extensions=[".docx",".pdf",".txt"],
                                 standard_names=["hm_clarification.docx","hm.docx"]),
    }

    for key, spec in asset_specs.items():
        path = find_one(**spec)
        try:
            text, refreshed = get_asset(key, path, manifest, force=args.force_refresh)
            out[f"{key}_text"] = text
            if path:
                out[f"{key}_file"] = path.name
            if refreshed:
                out["assets_refreshed"].append(key)
        except Exception as e:
            out["errors"].append(f"{key}: {e}")
            # Try reading stale cache as fallback
            cached = _read_cached_asset(key)
            if cached:
                out[f"{key}_text"] = cached
                out["errors"][-1] += " (using stale cache)"

    _save_manifest(manifest)

    # ── CVs (always fresh) ────────────────────────────────────────────────────
    for cv in find_cv_files():
        try:
            text = extract_file(cv)
            contact = extract_contact(text)
            out["cv_files"].append({
                "file_name": cv.name,
                "text": text,
                "email": contact["email"],
                "phone": contact["phone"],
            })
        except Exception as e:
            out["errors"].append(f"CV {cv.name}: {e}")

    # ── Print status to stderr, JSON to stdout ────────────────────────────────
    if out["assets_refreshed"]:
        print(f"[assets] Refreshed: {', '.join(out['assets_refreshed'])} → assets/",
              file=sys.stderr)
    else:
        print("[assets] All assets loaded from cache (assets/*.txt)", file=sys.stderr)

    print(json.dumps(out, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
